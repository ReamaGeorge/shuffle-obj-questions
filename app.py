from flask import Flask, render_template, request, send_file, url_for
from docx import Document
import random
import os
import tempfile

app = Flask(__name__)

# Temporary storage for documents
documents = {}

@app.route('/')
def index():
    return render_template('index.html', number_of_questions=5)  # Default to 5 questions

@app.route('/generate', methods=['POST'])
def generate():
    global documents  # Access the global documents dictionary

    number_of_questions = int(request.form['number_of_questions'])

    questions = []
    options = {}

    # Collect questions and options from the form
    for i in range(1, number_of_questions + 1):
        question = request.form[f'question_{i}'].capitalize()
        questions.append(question)

        options_for_question = {}
        for label in ['a', 'b', 'c', 'd']:
            option = request.form[f'option_{i}_{label}'].capitalize()
            options_for_question[label] = option

        options[question] = options_for_question

    # Generate 4 versions with shuffled questions and options
    versions = []
    for _ in range(4):
        shuffled_questions, shuffled_options = shuffle_questions(questions.copy(), options.copy())
        versions.append((shuffled_questions, shuffled_options))

    # Clear the documents dictionary
    documents.clear()

    # Create and save documents to temporary files
    for i, (questions, options) in enumerate(versions, start=1):
        doc = Document()
        for j, question in enumerate(questions, start=1):
            doc.add_paragraph(f"{j}. {question}")
            labels = ['a.', 'b.', 'c.', 'd.']
            for k, key in enumerate(options[question]):
                doc.add_paragraph(f"   {labels[k]} {options[question][key]}")
            doc.add_paragraph()

        # Save document to a temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(temp_file.name)
        temp_file.close()  # Close the file to allow downloading
        documents[f"Document {i}"] = temp_file.name  # Store the file path with a label

    # Generate download links for the documents
    download_links = [
        url_for('download', filename=label) for label in documents.keys()
    ]

    # Render the `generate` page with download links
    return render_template('download.html', download_links=download_links)

@app.route('/download/<filename>')
def download(filename):
    # Retrieve the temporary file path from the documents dictionary
    file_path = documents.get(filename)
    if file_path and os.path.exists(file_path):
        # Use call_on_close to delete the file after serving it
        def delete_file():
            try:
                os.remove(file_path)  # Delete the file after successful download
                documents.pop(filename, None)  # Remove the entry from the dictionary
            except Exception as e:
                app.logger.error(f"Error deleting file {file_path}: {e}")

        response = send_file(file_path, as_attachment=True)
        response.call_on_close(delete_file)
        return response
    else:
        return "File not found", 404

def shuffle_questions(questions, options):
    """
    Shuffle the list of questions and their corresponding options.
    """
    random.shuffle(questions)
    for question in questions:
        shuffled_options = options[question]
        shuffled_keys = list(shuffled_options.keys())
        random.shuffle(shuffled_keys)
        options[question] = {key: shuffled_options[key] for key in shuffled_keys}
    return questions, options

if __name__ == "__main__":
    app.run(debug=True)

